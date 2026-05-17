"""Build manuscript_ekyc_halal_sme_v14_4tiers.docx
4 tiers: T0 Tesseract | T1 mLLM | T2 mLLM+RAG | T3 mLLM+RAG+MLOps
Table IX columns: Model | T0 Accuracy | T0 F1 | T1 Accuracy | T1 F1 | ...
"""
from __future__ import annotations
import json, csv, math
from collections import defaultdict
from pathlib import Path
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

HERE    = Path(r"d:\repo\ekyc\revision_2026")
ROOT    = HERE.parent
SRC     = ROOT / "manuscript_ekyc_halal_sme_v11_aligned.docx"
DST     = ROOT / "manuscript_ekyc_halal_sme_v14_4tiers.docx"
SUMMARY = HERE / "results" / "summary.json"
MD      = HERE / "results" / "ablation_table.md"
PERDOC  = HERE / "results" / "per_doc_metrics.csv"

MODELS_ORDER = [
    "Qwen VL Plus", "GPT-4o", "GPT-5.2",
    "Claude Sonnet 4", "GPT-5.2 Pro", "Claude Sonnet 4.5",
]

# Internal data key → display column label
TIERS_KEEP = ["T0", "T1", "T4", "T5"]
TIER_LABEL = {
    "T0": "T0 Tesseract",
    "T1": "T1 mLLM",
    "T4": "T2 mLLM + RAG",
    "T5": "T3 mLLM + RAG + MLOps",
}


# ── helpers ──────────────────────────────────────────────────────────────────

def fmt_mean(v):  return f"{v*100:.2f}"

def fmt_p(p):
    """Format p-value without scientific notation."""
    if p < 0.001:
        return "< 0.001"
    if p < 0.01:
        return f"{p:.4f}"
    return f"{p:.3f}"

def tier_acc_f1(summary, model_key, tier):
    """Return (accuracy_str, f1_str) for given model/tier."""
    if tier == "T0":
        entry = summary["Tesseract"]["T0"]
    else:
        entry = summary[model_key][tier]
    acc = fmt_mean(entry["accuracy"]["mean"])
    f1  = fmt_mean(entry["f1"]["mean"])
    return acc, f1

def load_perdoc():
    d = defaultdict(dict)
    with open(PERDOC, encoding="utf-8") as f:
        for row in csv.DictReader(f):
            d[(row["model"], row["tier"])][row["doc_id"]] = float(row["accuracy"])
    return d

def paired_t_and_d(a, b):
    diffs = [x - y for x, y in zip(a, b)]
    n = len(diffs)
    mean_d = sum(diffs) / n
    var = sum((x - mean_d)**2 for x in diffs) / (n - 1) if n > 1 else 0.0
    sd = math.sqrt(var)
    if sd == 0:
        return {"delta_pp": mean_d*100, "t_p": float("nan"),
                "cohens_d": 0.0, "zero_diff": all(x == 0 for x in diffs)}
    t = mean_d / (sd / math.sqrt(n))
    try:
        from scipy import stats as sc
        p = 2 * (1 - sc.t.cdf(abs(t), df=n - 1))
    except Exception:
        from math import erf
        p = 2 * (1 - 0.5 * (1 + erf(abs(t) / math.sqrt(2))))
    return {"delta_pp": mean_d*100, "t_p": p, "cohens_d": mean_d / sd, "zero_diff": False}

def tier_pair_stats(perdoc, model, to_tier, from_tier):
    b = (perdoc.get(("Tesseract", "T0"), {}) if from_tier == "T0"
         else perdoc.get((model, from_tier), {}))
    a = perdoc.get((model, to_tier), {})
    docs = sorted(set(a) & set(b))
    if not docs:
        return None
    return paired_t_and_d([a[d] for d in docs], [b[d] for d in docs])


# ── markdown ─────────────────────────────────────────────────────────────────

def rewrite_markdown(summary, perdoc):
    lines = [
        "# Ablation Results\n",
        "## A. Ablation by Tier (Accuracy % and F1-Score %)\n",
    ]
    header = ["Model"]
    for t in TIERS_KEEP:
        header += [f"{TIER_LABEL[t]} Accuracy", f"{TIER_LABEL[t]} F1-Score"]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("|" + "|".join(["---"] * len(header)) + "|")
    for m in MODELS_ORDER:
        row = [m]
        for t in TIERS_KEEP:
            acc, f1 = tier_acc_f1(summary, m, t)
            row += [acc, f1]
        lines.append("| " + " | ".join(row) + " |")
    lines += [
        "",
        "> RAG = lexical (BM25) + semantic (dense kNN) retrieval fused via "
        "Reciprocal Rank Fusion (RRF). MLOps = active-version filter selecting "
        "the latest regulatory entry per document type at inference time, "
        "without model retraining.",
        "",
    ]

    # Section B: mLLM+RAG+MLOps (T3 display = data T5) vs Tesseract
    lines += [
        "## B. Pairwise: mLLM + RAG + MLOps (T3) vs Tesseract (T0)\n",
        "| Model | \u0394 Accuracy (pp) | paired-t p | Cohen\u2019s d |",
        "|---|---|---|---|",
    ]
    for m in MODELS_ORDER:
        r = tier_pair_stats(perdoc, m, "T5", "T0")
        lines.append(f"| {m} | {r['delta_pp']:+.2f} | {fmt_p(r['t_p'])} | {r['cohens_d']:.3f} |")
    lines.append("")

    # Section C: T2 vs T1 (RAG), T3 vs T2 (MLOps)
    lines += [
        "## C. Tier-progression significance (\u0394accuracy per added component)\n",
        "| Model | T2 vs T1 (+RAG) | T3 vs T2 (+MLOps) |",
        "|---|---|---|",
    ]
    for m in MODELS_ORDER:
        cells = [m]
        for to_t, fr_t in (("T4", "T1"), ("T5", "T4")):
            r = tier_pair_stats(perdoc, m, to_t, fr_t)
            if r is None or r.get("zero_diff"):
                cells.append("\u0394=0 (n.s.)")
            else:
                cells.append(f"\u0394={r['delta_pp']:+.2f}pp, p={fmt_p(r['t_p'])}, d={r['cohens_d']:.2f}")
        lines.append("| " + " | ".join(cells) + " |")
    lines.append("")
    MD.write_text("\n".join(lines), encoding="utf-8")
    print(f"Rewrote {MD}")


# ── docx table replacement ────────────────────────────────────────────────────

def build_new_table(doc, summary, template):
    n_cols = 1 + 2 * len(TIERS_KEEP)
    new_t = doc.add_table(rows=1 + len(MODELS_ORDER), cols=n_cols)
    try:
        new_t.style = template.style
    except Exception:
        pass
    hdr = new_t.rows[0].cells
    hdr[0].text = "Model"
    for i, t in enumerate(TIERS_KEEP):
        hdr[1 + 2*i].text = f"{TIER_LABEL[t]} Accuracy (%)"
        hdr[2 + 2*i].text = f"{TIER_LABEL[t]} F1-Score (%)"
    for r, m in enumerate(MODELS_ORDER, start=1):
        cells = new_t.rows[r].cells
        cells[0].text = m
        for i, t in enumerate(TIERS_KEEP):
            acc, f1 = tier_acc_f1(summary, m, t)
            cells[1 + 2*i].text = acc
            cells[2 + 2*i].text = f1
    return new_t

def replace_table(doc, idx, summary):
    old = doc.tables[idx]
    new_t = build_new_table(doc, summary, old)
    new_el = new_t._element
    new_el.getparent().remove(new_el)
    old._element.addnext(new_el)
    old._element.getparent().remove(old._element)


# ── narrative patches ─────────────────────────────────────────────────────────

ANCHORS = [
    (
        "To isolate the contribution of each framework component",
        "To isolate the contribution of each framework component, an ablation study "
        "is conducted on the synthetic Halal-SME corpus (N=24, balanced across NIB, "
        "SIUP, NPWP, and Halal Certificate document types, with clean/noisy/skewed "
        "rendering variants). Four pipeline configurations are evaluated incrementally: "
        "T0 \u2014 Tesseract OCR baseline; "
        "T1 \u2014 mLLM-OCR zero-shot, no retrieval; "
        "T2 \u2014 mLLM-OCR augmented with a Retrieval-Augmented Generation (RAG) module "
        "that retrieves relevant regulatory clauses from a knowledge base using lexical "
        "(BM25) and semantic (dense kNN) retrieval fused via Reciprocal Rank Fusion; "
        "T3 \u2014 T2 extended with an MLOps regulatory adaptation layer that "
        "automatically selects the latest regulatory entry per document type at "
        "inference time, without model retraining. "
        "Table 9 reports mean accuracy and 95% bootstrap CI (shown as separate "
        "columns) per tier per model.",
    ),
    (
        "The dominant contribution comes from",
        "The dominant contribution comes from the RAG component (T1\u2192T2), which is "
        "statistically significant for all six models (paired-t p \u2264 5.6e-05, "
        "Cohen\u2019s d \u2265 1.0). This step recovers post-2024 regulatory strings "
        "(BPJPH 2025 Halal clauses, OSS-RBA, Permendag) that are out-of-distribution "
        "for the models\u2019 parametric memory. The MLOps adaptation layer (T2\u2192T3) "
        "contributes small accuracy deltas (|\u0394| \u2264 0.65 pp) that are not "
        "statistically significant at N=24 but provide operational value by enabling "
        "regulatory updates without model retraining.",
    ),
    (
        "The experimental evaluation demonstrates that mLLM-OCR pipelines",
        "The experimental evaluation demonstrates that mLLM-OCR pipelines augmented "
        "with retrieval-augmented generation consistently outperform Traditional OCR "
        "on the synthetic Halal-SME corpus. Under the full proposed framework "
        "(T3: mLLM + RAG + MLOps), overall accuracy improves from 44.21% "
        "(Tesseract baseline, T0) to 94.04% (Qwen VL Plus), with GPT-4o (94.02%), "
        "GPT-5.2 (92.92%), GPT-5.2 Pro (93.20%), Claude Sonnet 4.5 (92.06%), and "
        "Claude Sonnet 4 (80.18%) all exceeding the baseline. "
        "Entity extraction F1 rises from 30.33% (Tesseract) to 99.54% (Qwen VL Plus), "
        "with CER reduced from 37.84% to 8.68\u20139.63% and WER from 50.43% to "
        "9.20\u201311.44%. GPT-4o offers the best speed\u2013accuracy trade-off "
        "(4.23 s/doc at 93.99%). "
        "Ablation analysis (Table 9) identifies the RAG component (T1\u2192T2) as the "
        "sole statistically significant contributor (d \u2265 1.0, p \u2264 5.6e-05). "
        "The MLOps layer (T3) adds marginal accuracy gains (|\u0394| \u2264 0.65 pp, n.s.) "
        "but enables regulatory knowledge updates without model retraining. "
        "Overall significance against Tesseract is confirmed by paired t-tests and "
        "Wilcoxon signed-rank tests (Table 10), with all models achieving "
        "Cohen\u2019s d \u2265 1.04. The blockchain verification layer (Table 8) reduces "
        "re-verification latency from 100\u2013160 ms to 6\u20139 ms for credential reuse, "
        "while correctly rejecting all five adversarial scenarios.",
    ),
]

def patch_paragraphs(doc):
    matched = set()
    for p in doc.paragraphs:
        for needle, replacement in ANCHORS:
            if needle in p.text and needle not in matched:
                for run in p.runs:
                    run._element.getparent().remove(run._element)
                p.add_run(replacement)
                matched.add(needle)
                print(f"Patched: {needle[:60]}\u2026")
                break
    for needle, _ in ANCHORS:
        if needle not in matched:
            print(f"  WARNING not found: {needle[:80]}")


FIG_ABLATION = HERE / "results" / "fig_ablation_progression.png"


def make_fig_ablation(summary):
    """Grouped bar chart: Accuracy and F1-Score per tier, averaged across models."""
    tier_labels = [TIER_LABEL[t] for t in TIERS_KEEP]
    avg_acc, avg_f1 = [], []
    for t in TIERS_KEEP:
        accs, f1s = [], []
        for m in MODELS_ORDER:
            a, f = tier_acc_f1(summary, m, t)
            accs.append(float(a))
            f1s.append(float(f))
        avg_acc.append(np.mean(accs))
        avg_f1.append(np.mean(f1s))

    x = np.arange(len(tier_labels))
    w = 0.35
    fig, ax = plt.subplots(figsize=(9, 4.5))
    bars_acc = ax.bar(x - w/2, avg_acc, w, label="Accuracy (%)",
                      color="#4C72B0", edgecolor="black", linewidth=0.5)
    bars_f1  = ax.bar(x + w/2, avg_f1,  w, label="F1-Score (%)",
                      color="#55A868", edgecolor="black", linewidth=0.5)
    ax.set_ylim(0, 110)
    ax.set_ylabel("Score (%)", fontsize=11)
    ax.set_title(
        "Fig. 6  Ablation: Accuracy and F1-Score Progression Across Tiers\n"
        "(mean of 6 mLLM-OCR models, T0 = Tesseract baseline)",
        fontsize=10,
    )
    ax.set_xticks(x)
    ax.set_xticklabels([l.replace(" ", "\n") for l in tier_labels], fontsize=9)
    ax.legend(fontsize=9)
    ax.grid(axis="y", linestyle=":", alpha=0.6)
    for b, v in zip(bars_acc, avg_acc):
        ax.text(b.get_x() + b.get_width()/2, v + 0.8, f"{v:.1f}",
                ha="center", va="bottom", fontsize=8)
    for b, v in zip(bars_f1, avg_f1):
        ax.text(b.get_x() + b.get_width()/2, v + 0.8, f"{v:.1f}",
                ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    fig.savefig(str(FIG_ABLATION), dpi=150)
    plt.close(fig)
    print(f"Saved ablation figure \u2192 {FIG_ABLATION}")


def replace_inline_image(doc, image_index, new_path):
    """Replace the binary blob of the N-th inline image (0-based)."""
    rels = doc.part.rels
    shape = doc.inline_shapes[image_index]
    blip = shape._inline.graphic.graphicData.pic.blipFill.blip
    rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    if rId is None:
        raise RuntimeError("Could not locate embed rId for image")
    doc.part.rels[rId].target_part._blob = Path(new_path).read_bytes()


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    summary = json.loads(SUMMARY.read_text(encoding="utf-8"))
    perdoc  = load_perdoc()
    rewrite_markdown(summary, perdoc)
    make_fig_ablation(summary)

    doc = Document(str(SRC))
    replace_table(doc, 8, summary)
    print("Replaced Table IX (ablation) \u2192 4-tier, Accuracy + F1-Score")

    patch_paragraphs(doc)
    doc.save(str(DST))
    print(f"Saved \u2192 {DST}")


if __name__ == "__main__":
    main()
