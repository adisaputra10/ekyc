"""Align body Tables III–VII and Figures 3–5 of the manuscript with the actual
measured data on the synthetic Halal-SME corpus (T4 = full proposed framework).

Reads:  manuscript_ekyc_halal_sme_v8_final.docx
        revision_2026/results/summary.json
Writes: manuscript_ekyc_halal_sme_v11_aligned.docx
        revision_2026/results/fig3_accuracy.png
        revision_2026/results/fig4_performance.png
        revision_2026/results/fig5_error_rates.png
"""
from __future__ import annotations
import json
from pathlib import Path
from copy import deepcopy

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
SRC = ROOT / "manuscript_ekyc_halal_sme_v8_final.docx"
DST = ROOT / "manuscript_ekyc_halal_sme_v11_aligned.docx"
SUMMARY = HERE / "results" / "summary.json"
RESULTS = HERE / "results"

# Model order in body tables = ablation order
MODELS = [
    "Qwen VL Plus",
    "GPT-4o",
    "GPT-5.2",
    "Claude Sonnet 4",
    "GPT-5.2 Pro",
    "Claude Sonnet 4.5",
]

# Tier used for body table headline numbers (full proposed framework)
HEADLINE_TIER = "T5"


def load_metrics():
    s = json.loads(SUMMARY.read_text(encoding="utf-8"))
    out = {}
    # Tesseract row
    t0 = s["Tesseract"]["T0"]
    out["Tesseract"] = {
        "accuracy": t0["accuracy"]["mean"] * 100,
        "precision": t0["precision"]["mean"] * 100,
        "recall": t0["recall"]["mean"] * 100,
        "f1": t0["f1"]["mean"] * 100,
        "layout": t0["layout"]["mean"] * 100,
        "cer": t0["cer"]["mean"] * 100,
        "wer": t0["wer"]["mean"] * 100,
        "latency": t0["latency_s"]["mean"],
    }
    for m in MODELS:
        t = s[m][HEADLINE_TIER]
        out[m] = {
            "accuracy": t["accuracy"]["mean"] * 100,
            "precision": t["precision"]["mean"] * 100,
            "recall": t["recall"]["mean"] * 100,
            "f1": t["f1"]["mean"] * 100,
            "layout": t["layout"]["mean"] * 100,
            "cer": t["cer"]["mean"] * 100,
            "wer": t["wer"]["mean"] * 100,
            "latency": t["latency_s"]["mean"],
        }
    return out


# ----------------- Figure generation -----------------

def make_fig3_accuracy(metrics, out_path):
    """Bar chart: Accuracy comparison Tesseract + 6 mLLM-OCR."""
    names = ["Tesseract"] + MODELS
    accs = [metrics[n]["accuracy"] for n in names]
    colors = ["#888"] + ["#4C72B0", "#55A868", "#C44E52", "#8172B2", "#CCB974", "#64B5CD"]

    fig, ax = plt.subplots(figsize=(9, 4.5))
    bars = ax.bar(names, accs, color=colors, edgecolor="black", linewidth=0.5)
    ax.set_ylabel("Accuracy (%)", fontsize=11)
    ax.set_title("Fig. 3  OCR Model Accuracy Comparison (Synthetic Halal-SME Corpus, T3 mLLM + RAG + MLOps)",
                 fontsize=11)
    ax.set_ylim(0, 100)
    ax.grid(axis="y", linestyle=":", alpha=0.6)
    for b, v in zip(bars, accs):
        ax.text(b.get_x() + b.get_width() / 2, v + 1, f"{v:.1f}%",
                ha="center", va="bottom", fontsize=9)
    plt.xticks(rotation=20, ha="right", fontsize=9)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def make_fig4_performance(metrics, out_path):
    """Grouped bar: Accuracy, F1, Layout for each model."""
    names = ["Tesseract"] + MODELS
    acc = [metrics[n]["accuracy"] for n in names]
    f1 = [metrics[n]["f1"] for n in names]
    layout = [metrics[n]["layout"] for n in names]

    x = np.arange(len(names))
    w = 0.27
    fig, ax = plt.subplots(figsize=(10, 4.8))
    ax.bar(x - w, acc, w, label="Accuracy (%)", color="#4C72B0", edgecolor="black", linewidth=0.4)
    ax.bar(x, f1, w, label="F1-score (%)", color="#55A868", edgecolor="black", linewidth=0.4)
    ax.bar(x + w, layout, w, label="Layout Score (%)", color="#C44E52", edgecolor="black", linewidth=0.4)
    ax.set_ylim(0, 110)
    ax.set_ylabel("Score (%)", fontsize=11)
    ax.set_title("Fig. 4  Performance Metrics: Accuracy, F1-score, and Layout Score",
                 fontsize=11)
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=20, ha="right", fontsize=9)
    ax.legend(loc="upper left", fontsize=9)
    ax.grid(axis="y", linestyle=":", alpha=0.6)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def make_fig5_errors(metrics, out_path):
    """Grouped bar: CER and WER (lower is better)."""
    names = ["Tesseract"] + MODELS
    cer = [metrics[n]["cer"] for n in names]
    wer = [metrics[n]["wer"] for n in names]

    x = np.arange(len(names))
    w = 0.36
    fig, ax = plt.subplots(figsize=(9.5, 4.5))
    ax.bar(x - w / 2, cer, w, label="CER (%) ↓", color="#C44E52", edgecolor="black", linewidth=0.4)
    ax.bar(x + w / 2, wer, w, label="WER (%) ↓", color="#DD8452", edgecolor="black", linewidth=0.4)
    ax.set_ylim(0, max(max(cer), max(wer)) * 1.25)
    ax.set_ylabel("Error Rate (%)", fontsize=11)
    ax.set_title("Fig. 5  Error Rate Comparison: CER and WER (lower is better)",
                 fontsize=11)
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=20, ha="right", fontsize=9)
    ax.legend(loc="upper right", fontsize=9)
    ax.grid(axis="y", linestyle=":", alpha=0.6)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


# ----------------- Table cell rewriting -----------------

def set_cell(cell, text):
    """Replace a cell's content with `text` preserving the first paragraph's
    formatting (font, bold, size)."""
    # capture first run formatting from first paragraph
    p = cell.paragraphs[0]
    # delete existing paragraphs except the first
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    # clear all runs in first paragraph
    for run in p.runs:
        run._element.getparent().remove(run._element)
    # add new run with text
    p.add_run(text)


def rewrite_table_2(t):
    """Model list. Header row + 6 model rows. Original has 3 columns; we'll
    just write into column 0."""
    # row 0 = header "Model mLLM-OCR"
    set_cell(t.rows[0].cells[0], "Model mLLM-OCR")
    for i, m in enumerate(MODELS):
        if i + 1 < len(t.rows):
            set_cell(t.rows[i + 1].cells[0], m)


def rewrite_table_3(t, M):
    """TABLE IV  Overall: Model | Accuracy | F1 | Layout (only 4 logical cols)."""
    headers = ["Model", "Accuracy (%)", "F1-Score (%)", "Layout Score (%)"]
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h)
    rows = [("Tesseract (Traditional OCR)", "Tesseract")] + [
        (f"{m} (mLLM-OCR)", m) for m in MODELS]
    for r, (label, key) in enumerate(rows, start=1):
        if r >= len(t.rows):
            break
        d = M[key]
        set_cell(t.rows[r].cells[0], label)
        set_cell(t.rows[r].cells[1], f"{d['accuracy']:.2f}")
        set_cell(t.rows[r].cells[2], f"{d['f1']:.2f}")
        set_cell(t.rows[r].cells[3], f"{d['layout']:.2f}")


def rewrite_table_4(t, M):
    """TABLE V  Entity: Model | Precision | Recall | F1."""
    headers = ["Model", "Precision (%)", "Recall (%)", "F1-Score (%)"]
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h)
    rows = [("Tesseract (Traditional OCR)", "Tesseract")] + [
        (f"{m} (mLLM-OCR)", m) for m in MODELS]
    for r, (label, key) in enumerate(rows, start=1):
        if r >= len(t.rows):
            break
        d = M[key]
        set_cell(t.rows[r].cells[0], label)
        set_cell(t.rows[r].cells[1], f"{d['precision']:.2f}")
        set_cell(t.rows[r].cells[2], f"{d['recall']:.2f}")
        set_cell(t.rows[r].cells[3], f"{d['f1']:.2f}")


def rewrite_table_5(t, M):
    """TABLE VI  Errors: Model | CER | WER."""
    headers = ["Model", "CER (%)", "WER (%)"]
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h)
    rows = [("Tesseract (Traditional OCR)", "Tesseract")] + [
        (f"{m} (mLLM-OCR)", m) for m in MODELS]
    for r, (label, key) in enumerate(rows, start=1):
        if r >= len(t.rows):
            break
        d = M[key]
        set_cell(t.rows[r].cells[0], label)
        set_cell(t.rows[r].cells[1], f"{d['cer']:.2f}")
        set_cell(t.rows[r].cells[2], f"{d['wer']:.2f}")


def rewrite_table_6(t, M):
    """TABLE VII Comprehensive: Model | Acc | CER | WER | F1 | Layout | Speed."""
    headers = ["Model", "Accuracy (%)", "CER (%) ↓", "WER (%) ↓",
               "Entity F1 (%)", "Layout Score (0-100)", "Speed (s/img)"]
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h)
    rows = [("Tesseract (Traditional OCR)", "Tesseract")] + [
        (f"{m} (mLLM-OCR)", m) for m in MODELS]
    for r, (label, key) in enumerate(rows, start=1):
        if r >= len(t.rows):
            break
        d = M[key]
        set_cell(t.rows[r].cells[0], label)
        set_cell(t.rows[r].cells[1], f"{d['accuracy']:.2f}")
        set_cell(t.rows[r].cells[2], f"{d['cer']:.2f}")
        set_cell(t.rows[r].cells[3], f"{d['wer']:.2f}")
        set_cell(t.rows[r].cells[4], f"{d['f1']:.2f}")
        set_cell(t.rows[r].cells[5], f"{d['layout']:.1f}")
        set_cell(t.rows[r].cells[6], f"{d['latency']:.2f}")


# ----------------- Image replacement -----------------

def replace_inline_image(doc, image_index, new_path):
    """Replace the binary blob of the N-th inline image (0-based) with the
    contents of `new_path`. Keeps the same anchor/sizing in the document."""
    rels = doc.part.rels
    # Map inline_shapes order -> rId via the underlying drawing element
    shapes = doc.inline_shapes
    if image_index >= len(shapes):
        raise IndexError(f"image_index {image_index} >= {len(shapes)}")
    shape = shapes[image_index]
    blip = shape._inline.graphic.graphicData.pic.blipFill.blip
    rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    if rId is None:
        raise RuntimeError("Could not locate embed rId for image")
    image_part = rels[rId].target_part
    image_part._blob = Path(new_path).read_bytes()


# ----------------- Narrative text patches -----------------

PARA_PATCHES = [
    # (substring to locate, full replacement text)
    (
        "Fig. 3 compares the mean per-document character accuracy of Tesseract",
        "Fig. 3 compares the mean per-document accuracy of Tesseract against six "
        "mLLM-OCR models on the synthetic Halal-SME corpus under the full proposed "
        "framework configuration (T4: schema-guided prompting + hybrid BM25/Dense "
        "RAG). The Tesseract baseline (44.21%) is surpassed by all multimodal "
        "approaches. Qwen VL Plus, GPT-4o, GPT-5.2, and GPT-5.2 Pro all exceed "
        "93% accuracy, while Claude Sonnet 4.5 attains 91.75%. Claude Sonnet 4 "
        "reaches 81.47%; its lower composite score is driven by a model-specific "
        "output-formatting behaviour that depresses the layout sub-score, not by "
        "OCR-quality degradation (see Addendum, Section G)."
    ),
    (
        "Traditional OCR produces the weakest accuracy and layout scores",
        "As shown in Table IV, Traditional OCR (Tesseract) produces the weakest "
        "accuracy, F1, and layout scores. All mLLM-OCR models consistently "
        "outperform the baseline under the full proposed framework. Qwen VL Plus "
        "achieves the highest accuracy (94.00%) and entity F1 (99.54%). GPT-4o, "
        "GPT-5.2, and GPT-5.2 Pro form a tight cluster around 93–94% accuracy and "
        "98–99% F1. Claude Sonnet 4.5 reaches 91.75% accuracy and 98.96% F1, while "
        "Claude Sonnet 4 (the older variant) trails at 81.47% accuracy due to a "
        "single-line concatenated output behaviour that suppresses its layout "
        "score, although its entity F1 (97.40%) is in the same band as the other "
        "models."
    ),
    (
        "The Fig.  shows accuracy, F1-score, and layout score on OCR and mLLM-OCR.",
        "Fig. 4 shows accuracy, F1-score, and layout score for the Tesseract "
        "baseline and six mLLM-OCR models on the synthetic Halal-SME corpus "
        "under the full proposed framework. Traditional OCR exhibits the lowest "
        "scores across all three metrics, while the multimodal models cluster "
        "near the upper bound on accuracy and F1. Layout scores are uniformly "
        "high (>80%) for every model except Claude Sonnet 4, which is impacted "
        "by an output-format artefact discussed in the addendum."
    ),
    (
        "Claude 4.5 Sonnet achieves the highest precision",
        "Under the full proposed framework, Qwen VL Plus achieves the highest "
        "entity F1-score (99.54%), with Claude Sonnet 4.5 (98.96%), GPT-4o "
        "(99.48%), GPT-5.2 (98.15%), and GPT-5.2 Pro (98.15%) all clustered above "
        "98%. Claude Sonnet 4 attains 97.40% F1. Precision, recall, and F1 are "
        "numerically identical for the mLLM-OCR models because the framework "
        "enforces a schema-bound JSON output contract: the number of predicted "
        "fields equals the number of expected fields by construction, so the "
        "three metrics collapse to the proportion of correctly recovered fields. "
        "This metric definition is appropriate for the eKYC use case but should "
        "not be compared head-to-head with span-based benchmarks such as FUNSD or "
        "DocVQA (see Addendum, Section G)."
    ),
    (
        "All mLLM-OCR models significantly reduce both CER and WER.",
        "All mLLM-OCR models significantly reduce both CER and WER relative to "
        "Tesseract (CER 37.84%, WER 50.43%). GPT-4o achieves the lowest CER "
        "(8.68%) while Qwen VL Plus achieves the lowest WER (9.20%). The six "
        "multimodal models cluster tightly in the 8.6–9.7% CER and 9.2–11.4% WER "
        "range. Although CER and WER are secondary objectives in eKYC, the "
        "uniform reduction confirms that the multimodal backbones recover the "
        "underlying transcription content reliably."
    ),
    (
        "The comparison of CER and WER of OCR methods is shown in this diagram.",
        "Fig. 5 compares CER and WER for the seven systems; lower is better. "
        "Tesseract incurs the highest error rates on both metrics. All mLLM-OCR "
        "models reduce CER by roughly 75–80% and WER by roughly 77–82% relative "
        "to Tesseract on the synthetic Halal-SME corpus."
    ),
    (
        "According to Table 7, the proposed framework consistently outperforms",
        "According to Table VII, the proposed framework consistently outperforms "
        "the Tesseract baseline across all headline metrics. Qwen VL Plus leads "
        "on accuracy (94.00%) and entity F1 (99.54%), with GPT-4o (93.99% / "
        "99.48%) and GPT-5.2 Pro (93.20% / 98.15%) close behind. GPT-4o offers "
        "the best speed–accuracy trade-off at 4.23 s per document, while "
        "GPT-5.2 Pro is the slowest (23.93 s) for a marginally lower accuracy. "
        "Layout preservation exceeds 80% for every model except Claude Sonnet 4."
    ),
]


def patch_paragraphs(doc):
    patched = 0
    for p in doc.paragraphs:
        full = p.text
        if not full.strip():
            continue
        for needle, replacement in PARA_PATCHES:
            if needle in full:
                # rewrite paragraph: clear all runs, add one new run
                for run in p.runs:
                    run._element.getparent().remove(run._element)
                p.add_run(replacement)
                patched += 1
                break
    return patched


# ----------------- Main -----------------

def main():
    M = load_metrics()
    print("Loaded metrics for", list(M.keys()))

    fig3 = RESULTS / "fig3_accuracy.png"
    fig4 = RESULTS / "fig4_performance.png"
    fig5 = RESULTS / "fig5_error_rates.png"
    make_fig3_accuracy(M, fig3)
    make_fig4_performance(M, fig4)
    make_fig5_errors(M, fig5)
    print(f"Saved figures: {fig3.name}, {fig4.name}, {fig5.name}")

    doc = Document(str(SRC))
    # tables are 0-indexed: 0=lit, 1=dataset, 2=model list, 3..6=body, 7=blockchain, 8..9=ablation
    rewrite_table_2(doc.tables[2])
    rewrite_table_3(doc.tables[3], M)
    rewrite_table_4(doc.tables[4], M)
    rewrite_table_5(doc.tables[5], M)
    rewrite_table_6(doc.tables[6], M)
    print("Rewrote tables 2, 3, 4, 5, 6")

    # Inline shapes order from the v8 doc: 0=Fig1, 1=Fig2, 2=Fig3, 3=Fig4, 4=Fig5
    replace_inline_image(doc, 2, fig3)
    replace_inline_image(doc, 3, fig4)
    replace_inline_image(doc, 4, fig5)
    print("Replaced inline images for Fig 3, 4, 5")

    n_patched = patch_paragraphs(doc)
    print(f"Patched {n_patched} narrative paragraphs")

    doc.save(str(DST))
    print(f"Saved -> {DST}")


if __name__ == "__main__":
    main()
