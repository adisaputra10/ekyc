"""Drop the T2 (+schema) column from the ablation table, split each cell into
separate Mean and 95% CI columns, and refresh the narrative.

Reads:  manuscript_ekyc_halal_sme_v11_aligned.docx
        revision_2026/results/summary.json
Writes: manuscript_ekyc_halal_sme_v12_no_t2.docx
        revision_2026/results/ablation_table.md (updated in place)
"""
from __future__ import annotations
import json
import csv
import math
import statistics as pystats
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
SRC = ROOT / "manuscript_ekyc_halal_sme_v11_aligned.docx"
DST = ROOT / "manuscript_ekyc_halal_sme_v12_no_t2.docx"
SUMMARY = HERE / "results" / "summary.json"
MD = HERE / "results" / "ablation_table.md"
PERDOC = HERE / "results" / "per_doc_metrics.csv"
STATS = HERE / "results" / "statistics.json"

MODELS_ORDER = [
    "Qwen VL Plus",
    "GPT-4o",
    "GPT-5.2",
    "Claude Sonnet 4",
    "GPT-5.2 Pro",
    "Claude Sonnet 4.5",
]
TIERS_KEEP = ["T0", "T1", "T3", "T4", "T5"]
TIER_LABEL = {
    "T0": "T0 Tesseract",
    "T1": "T1 mLLM zero-shot",
    "T3": "T2 +Dense RAG",
    "T4": "T3 +Hybrid RAG",
    "T5": "T4 +MLOps",
}


def fmt_mean(v):
    return f"{v * 100:.2f}"


def fmt_ci(lo, hi):
    return f"{lo * 100:.2f} \u2013 {hi * 100:.2f}"


def load_perdoc():
    """Return per_doc[(model, tier)] -> dict[doc_id -> accuracy]."""
    d = defaultdict(dict)
    with open(PERDOC, encoding="utf-8") as f:
        for row in csv.DictReader(f):
            d[(row["model"], row["tier"])][row["doc_id"]] = float(row["accuracy"])
    return d


def paired_t_and_d(a, b):
    """Paired t-test p-value and Cohen's d (paired) for accuracy a vs b on same docs."""
    diffs = [x - y for x, y in zip(a, b)]
    n = len(diffs)
    mean_d = sum(diffs) / n
    var = sum((x - mean_d) ** 2 for x in diffs) / (n - 1) if n > 1 else 0.0
    sd = math.sqrt(var)
    if sd == 0:
        return {"delta_pp": mean_d * 100, "t_p": float("nan"), "cohens_d": 0.0,
                "zero_diff": all(x == 0 for x in diffs)}
    t = mean_d / (sd / math.sqrt(n))
    # two-sided p via Student-t survival; use scipy if present, else approximate
    try:
        from scipy import stats as scistats
        p = 2 * (1 - scistats.t.cdf(abs(t), df=n - 1))
    except Exception:
        # normal approximation fallback
        from math import erf
        p = 2 * (1 - 0.5 * (1 + erf(abs(t) / math.sqrt(2))))
    d = mean_d / sd
    return {"delta_pp": mean_d * 100, "t_p": p, "cohens_d": d, "zero_diff": False}


def compute_tier_pair(perdoc, model, tier_to, tier_from):
    """Compute paired stats for (model, tier_to) - (model, tier_from)."""
    a = perdoc.get((model, tier_to), {})
    b = perdoc.get((model, tier_from), {})
    docs = sorted(set(a) & set(b))
    if not docs:
        return None
    av = [a[d] for d in docs]
    bv = [b[d] for d in docs]
    return paired_t_and_d(av, bv)


def tier_value(summary, model_key, tier):
    """Return (mean_pct_str, ci_str) for a (model, tier). For T0 every model
    shares the Tesseract baseline."""
    if tier == "T0":
        td = summary["Tesseract"]["T0"]["accuracy"]
    else:
        td = summary[model_key][tier]["accuracy"]
    return fmt_mean(td["mean"]), fmt_ci(td["ci_low"], td["ci_high"])


# ---------- Markdown rewrite ----------

def rewrite_markdown(summary, perdoc):
    lines = []
    lines.append("# Ablation Results\n")
    lines.append("## A. Ablation by Tier (accuracy %, mean and 95% bootstrap CI)\n")
    # build header
    header_cells = ["Model"]
    for t in TIERS_KEEP:
        header_cells.append(f"{TIER_LABEL[t]} Mean")
        header_cells.append(f"{TIER_LABEL[t]} 95% CI")
    lines.append("| " + " | ".join(header_cells) + " |")
    lines.append("|" + "|".join(["---"] * len(header_cells)) + "|")
    for m in MODELS_ORDER:
        row = [m]
        for t in TIERS_KEEP:
            mean, ci = tier_value(summary, m, t)
            row.append(mean)
            row.append(ci)
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    lines.append("> Note: T2 (+schema only) is omitted from the headline table; "
                 "see Section C and the Summary for the schema-only ablation result.")
    lines.append("")

    # Section B — recomputed from per-doc CSV to be self-consistent
    lines.append("## B. Pairwise: each mLLM (best tier T4) vs Tesseract (T0)\n")
    lines.append("| Model | Δ Accuracy (pp) | paired-t p | Cohen's d |")
    lines.append("|---|---|---|---|")
    for m in MODELS_ORDER:
        r = compute_tier_pair(perdoc, m, "T4", "T0")
        # Tesseract entries for T0 only live under 'Tesseract' model name in CSV;
        # build the comparison vector manually:
        a = perdoc.get((m, "T4"), {})
        b = perdoc.get(("Tesseract", "T0"), {})
        docs = sorted(set(a) & set(b))
        av = [a[d] for d in docs]
        bv = [b[d] for d in docs]
        r = paired_t_and_d(av, bv)
        lines.append(f"| {m} | {r['delta_pp']:+.2f} | {r['t_p']:.2e} | {r['cohens_d']:.3f} |")
    lines.append("")

    # Section C — T2 removed; compute T3 vs T1, T4 vs T3, T5 vs T4 directly
    lines.append("## C. Tier-progression significance (Δaccuracy per added component)\n")
    lines.append("| Model | T2 vs T1 (Dense RAG) | T3 vs T2 (Hybrid) | T4 vs T3 (MLOps) |")
    lines.append("|---|---|---|---|")
    for m in MODELS_ORDER:
        cells = [m]
        for to_tier, from_tier in (("T3", "T1"), ("T4", "T3"), ("T5", "T4")):
            r = compute_tier_pair(perdoc, m, to_tier, from_tier)
            if r is None or r.get("zero_diff"):
                cells.append("Δ=0 (n.s.)")
            else:
                cells.append(f"Δ={r['delta_pp']:+.2f}pp, p={r['t_p']:.2e}, d={r['cohens_d']:.2f}")
        lines.append("| " + " | ".join(cells) + " |")
    lines.append("")
    lines.append("> Note: T2 (+schema) was removed from the headline tier sequence because "
                 "the schema-only step produced |\u0394| \u2264 0.74 pp across all models, none "
                 "statistically significant at N=24 (all p > 0.32). The remaining tiers are "
                 "relabelled T0\u2013T4 for continuity: T0 Tesseract, T1 mLLM zero-shot, "
                 "T2 +Dense RAG, T3 +Hybrid RAG, T4 +MLOps. Schema-guided JSON output is "
                 "still retained in the operational pipeline for its engineering value "
                 "(parsability, validation, downstream contract) \u2014 see Section G.")
    lines.append("")
    MD.write_text("\n".join(lines), encoding="utf-8")
    print(f"Rewrote {MD}")


# ---------- DOCX table replacement ----------

def build_new_table(doc, summary, template_table):
    """Build a new docx table with 11 columns replacing template_table."""
    n_rows = 1 + len(MODELS_ORDER)
    n_cols = 1 + 2 * len(TIERS_KEEP)
    new_t = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        new_t.style = template_table.style
    except Exception:
        pass

    # Header
    hdr = new_t.rows[0].cells
    hdr[0].text = "Model"
    for i, t in enumerate(TIERS_KEEP):
        hdr[1 + 2 * i].text = f"{TIER_LABEL[t]} Mean"
        hdr[2 + 2 * i].text = f"{TIER_LABEL[t]} 95% CI"

    # Data rows
    for r, m in enumerate(MODELS_ORDER, start=1):
        cells = new_t.rows[r].cells
        cells[0].text = m
        for i, t in enumerate(TIERS_KEEP):
            mean, ci = tier_value(summary, m, t)
            cells[1 + 2 * i].text = mean
            cells[2 + 2 * i].text = ci
    return new_t


def replace_table(doc, table_index, summary):
    """Replace doc.tables[table_index] with the new 11-column structure."""
    old = doc.tables[table_index]
    new_t = build_new_table(doc, summary, old)
    # Move new_t._element to right after old._element and remove old
    old_el = old._element
    new_el = new_t._element
    # `add_table` appended new_el to body; cut it out and insert after old_el
    new_el.getparent().remove(new_el)
    old_el.addnext(new_el)
    old_el.getparent().remove(old_el)


# ---------- Paragraph patches ----------

NEW_PARA_99 = (
    "To isolate the contribution of each framework component, an ablation study "
    "is conducted on the synthetic Halal-SME corpus (N=24, balanced across NIB, "
    "SIUP, NPWP, and Halal Certificate document types, with clean/noisy/skewed "
    "rendering variants). Five pipeline configurations are reported: "
    "T0 — Tesseract baseline; T1 — mLLM-OCR in zero-shot mode; "
    "T3 — T1 augmented with Dense kNN RAG retrieval over the regulatory "
    "knowledge base; T4 — T3 with Hybrid retrieval (BM25 + Dense via Reciprocal "
    "Rank Fusion); T5 — T4 with MLOps regulatory adaptation. A schema-only "
    "configuration (T2 = T1 + schema-guided prompting) was also evaluated but "
    "produced negligible accuracy change (|Δ| ≤ 0.74 pp, p > 0.32 for all six "
    "models) and is therefore omitted from the headline table; the schema "
    "contract is nonetheless retained in the operational pipeline for "
    "engineering reasons (structured JSON output, validation, downstream "
    "contract). Table 9 reports mean accuracy with 95% bootstrap confidence "
    "intervals per tier per model (Mean and 95% CI shown in separate columns "
    "for clarity)."
)

NEW_PARA_106 = (
    "The dominant contribution comes from Dense RAG (T1\u2192T2), which is "
    "statistically significant for all six models (paired-t p \u2264 5.6e-05, "
    "Cohen\u2019s d \u2265 1.0). This step recovers post-2024 regulatory strings (BPJPH "
    "2025 Halal clauses, OSS-RBA, Permendag) that are out-of-distribution for "
    "the models\u2019 parametric memory. Hybrid retrieval (T2\u2192T3) and MLOps "
    "adaptation (T3\u2192T4) each contribute small accuracy deltas (|\u0394| \u2264 0.65 pp) "
    "that are not individually statistically significant at N=24 but provide "
    "operational value: hybrid retrieval improves robustness for queries with "
    "regulatory number anchors, and MLOps adaptation eliminates version "
    "ambiguity during regulatory updates without model retraining."
)

NEW_PARA_113 = (
    "The experimental evaluation demonstrates that mLLM-OCR pipelines, when "
    "combined with retrieval-augmented generation over a regulatory knowledge "
    "base, consistently outperform Traditional OCR on the synthetic Halal-SME "
    "corpus. Under the full proposed framework (T4) overall accuracy improves "
    "from 44.21% (Tesseract baseline) to 94.00% (Qwen VL Plus, best), with "
    "GPT-4o (93.99%), GPT-5.2 (93.20%), GPT-5.2 Pro (93.20%), Claude Sonnet "
    "4.5 (91.75%), and Claude Sonnet 4 (81.47%) clustered above the baseline. "
    "Entity extraction F1 rises from 30.33% (Tesseract) to 99.54% (Qwen VL "
    "Plus), accompanied by substantial reductions in CER (37.84% → 8.68–9.63%) "
    "and WER (50.43% → 9.20–11.44%). GPT-4o offers the most favourable "
    "speed–accuracy trade-off (4.23 s per document at 93.99% accuracy). "
    "Ablation analysis (Table 9) identifies Dense RAG retrieval (T1\u2192T2) as the "
    "dominant contributing component, producing large and statistically "
    "significant accuracy uplifts (d \u2265 1.0, p \u2264 5.6e-05) across all six models. "
    "The full proposed framework (T3: Hybrid RAG) and MLOps adaptation (T4) add "
    "marginal gains (|\u0394| \u2264 0.65 pp, all n.s.) but deliver operational value in "
    "robustness and regulatory-update agility. Overall significance against "
    "Tesseract is confirmed by paired t-tests and Wilcoxon signed-rank tests "
    "(Table 10), with all models achieving Cohen\u2019s d \u2265 1.04 (large effect). The "
    "blockchain verification layer (Table 8) reduces re-verification latency "
    "from 100\u2013160 ms during initial AI inference to 6\u20139 ms for credential reuse, "
    "while correctly rejecting all five adversarial scenarios tested, providing "
    "a complete and auditable identity management chain for Halal SME onboarding."
)


def replace_paragraph_text(p, new_text):
    for run in p.runs:
        run._element.getparent().remove(run._element)
    p.add_run(new_text)


def patch_paragraphs(doc):
    # Anchor by substring to be resilient to index shifts
    anchors = [
        ("To isolate the contribution of each framework component", NEW_PARA_99),
        ("The dominant contribution comes from the Dense RAG step", NEW_PARA_106),
        ("The experimental evaluation demonstrates that mLLM-OCR pipelines",
         NEW_PARA_113),
    ]
    matched = set()
    for p in doc.paragraphs:
        for needle, replacement in anchors:
            if needle in p.text and needle not in matched:
                replace_paragraph_text(p, replacement)
                matched.add(needle)
                print(f"Patched: {needle[:60]}…")
                break
    if len(matched) != len(anchors):
        missing = [n for n, _ in anchors if n not in matched]
        for m in missing:
            print(f"  WARNING: anchor not found: {m[:80]}")


def main():
    summary = json.loads(SUMMARY.read_text(encoding="utf-8"))
    perdoc = load_perdoc()
    rewrite_markdown(summary, perdoc)

    doc = Document(str(SRC))
    # Ablation table is doc.tables[8] (TABLE IX in caption)
    replace_table(doc, 8, summary)
    print("Replaced Table 8 (Ablation) with 11-column Mean/CI split, T2 removed")

    patch_paragraphs(doc)
    DST_NEW = ROOT / "manuscript_ekyc_halal_sme_v13_tiers.docx"
    doc.save(str(DST_NEW))
    print(f"Saved -> {DST_NEW}")


if __name__ == "__main__":
    main()
